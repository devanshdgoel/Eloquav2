from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_paragraph(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def add_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_callout(doc, text, bg_hex='E8F5F4'):
    p = doc.add_paragraph()
    shade_paragraph(p, bg_hex)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '240')
    ind.set(qn('w:right'), '240')
    pPr.append(ind)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def add_screenshot_placeholder(doc, label, height_inches=3.8):
    """
    Insert a centred, bordered grey box to hold a phone screenshot.
    Width is fixed at 2.5 inches (portrait phone frame).
    """
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tbl.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()

    # Cell width: 2.5 inches in twips (1 inch = 1440 twips)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(2.5 * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

    # Row height
    tr = cell._tc.getparent()
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(height_inches * 1440)))
    trH.set(qn('w:hRule'), 'exact')
    trPr.append(trH)

    # Light grey fill
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    tcPr.append(shd)

    # Thin grey border on all four sides
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement('w:' + side)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'AAAAAA')
        tcBorders.append(el)
    tcPr.append(tcBorders)

    # Label text, pushed toward the vertical centre with top spacing
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(height_inches * 36 - 20)

    cam_run = p.add_run('[  ')
    cam_run.font.size = Pt(11)
    cam_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    label_run = p.add_run('Screenshot: ' + label)
    label_run.font.size = Pt(10)
    label_run.font.italic = True
    label_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    end_run = p.add_run('  ]')
    end_run.font.size = Pt(11)
    end_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_paragraph()


def build_doc(platform):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(13)

    is_ios = (platform == 'ios')

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run('Eloqua')
    tr.bold = True
    tr.font.size = Pt(32)
    tr.font.color.rgb = RGBColor(0x2B, 0x7A, 0x78)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = sub_p.add_run(
        'Getting Started  -  iPhone Guide' if is_ios else
        'Getting Started  -  Android Guide'
    )
    sub.font.size = Pt(17)
    sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    add_callout(
        doc,
        'This guide will help you install and open Eloqua for the very first time.\n'
        'Each step is short. Take your time. There is no rush.',
        bg_hex='E8F5F4'
    )
    doc.add_paragraph()

    # What you need
    h = doc.add_paragraph()
    hr = h.add_run('What you will need')
    hr.bold = True
    hr.font.size = Pt(16)
    hr.font.color.rgb = RGBColor(0x2B, 0x7A, 0x78)

    needs = (
        [
            'An iPhone running iOS 16 or newer',
            'A Wi-Fi or mobile data connection',
            'The invitation email from Eloqua (check your inbox and spam folder)',
        ] if is_ios else [
            'An Android phone running Android 10 or newer',
            'A Wi-Fi or mobile data connection',
            'A Google account (the one you use for the Play Store on your phone)',
        ]
    )
    for item in needs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item).font.size = Pt(13)
        p.paragraph_format.left_indent = Inches(0.3)

    doc.add_paragraph()
    add_rule(doc)
    doc.add_paragraph()

    # Step helpers
    def step_heading(number, title):
        p = doc.add_paragraph()
        r1 = p.add_run('Step {}:  '.format(number))
        r1.bold = True
        r1.font.size = Pt(15)
        r1.font.color.rgb = RGBColor(0xFF, 0xA9, 0x40)
        r2 = p.add_run(title)
        r2.bold = True
        r2.font.size = Pt(15)
        r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(3)

    def body(text):
        p = doc.add_paragraph()
        p.add_run(text).font.size = Pt(13)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    def tip(text):
        p = doc.add_paragraph()
        shade_paragraph(p, 'FFF8EE')
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '200')
        pPr.append(ind)
        r = p.add_run('Tip:  ' + text)
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x5C, 0x3D, 0x00)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)

    def ss(label, h=3.8):
        add_screenshot_placeholder(doc, label, height_inches=h)

    # iOS steps
    if is_ios:

        step_heading(1, 'Find the invitation email')
        body('Open the email app on your iPhone.')
        body('Look for an email from Eloqua or TestFlight.')
        body('The subject line will say "You have been invited to test Eloqua".')
        tip('If you cannot find it, check your Junk or Spam folder.')
        ss('Invitation email showing the "Start testing" link', h=2.5)

        step_heading(2, 'Download the TestFlight app')
        body('TestFlight is a free Apple app that lets you try apps before they are in the App Store.')
        body('Open the App Store on your iPhone.')
        body('Tap the Search icon at the bottom of the screen.')
        body('Type  TestFlight  and tap Search.')
        body('Tap Get next to the TestFlight app (it has an orange and white icon).')
        body('Wait for it to finish downloading, then tap Open.')
        tip('If TestFlight is already on your phone, skip to Step 3.')
        ss('TestFlight in the App Store showing the Get button', h=3.8)

        step_heading(3, 'Accept the invitation')
        body('Go back to the invitation email.')
        body('Tap the blue link that says  Start testing  or  View in TestFlight.')
        body('Your iPhone will open TestFlight and show the Eloqua app.')
        body('Tap  Accept.')
        ss('TestFlight showing the Eloqua invitation with the Accept button', h=3.8)

        step_heading(4, 'Install Eloqua')
        body('On the Eloqua page in TestFlight, tap  Install.')
        body('Eloqua will download to your phone. This usually takes less than a minute.')
        body('When the download is done, tap  Open.')
        tip('You can also find the Eloqua icon on your home screen and tap it there.')
        ss('Eloqua listed in TestFlight with the Install button', h=3.8)

        step_heading(5, 'Create your account')
        body('Tap  Sign Up.')
        body('Enter your email address.')
        body('Choose a password. Use at least 8 characters.')
        body('Tap  Create Account.')
        tip('Write your email and password down somewhere safe before continuing.')
        ss('Eloqua Sign Up screen', h=3.8)

        step_heading(6, 'Allow microphone access')
        body('A message will appear: "Eloqua would like to access the Microphone".')
        body('Tap  Allow.')
        body('The microphone is needed so Eloqua can hear your voice during exercises.')
        tip(
            'If you tapped  Do Not Allow  by mistake, go to:\n'
            'Settings  >  Privacy & Security  >  Microphone\n'
            'Find Eloqua in the list and turn it on.'
        )
        ss('iPhone microphone permission popup', h=2.5)

        step_heading(7, 'Tell us about yourself')
        body('Type your first name.')
        body('Enter your age.')
        body('Tap  Continue.')
        ss('Eloqua About You screen', h=3.8)

        step_heading(8, 'Record your voice  (3 short sentences)')
        body('You will be asked to read 3 sentences out loud.')
        body('This helps Eloqua learn to recognise your voice.')
        body('Tap the record button, read the sentence clearly, then tap stop.')
        body('Repeat this for all 3 sentences.')
        tip('Find a quiet room. Hold your phone about 20 cm (8 inches) from your mouth.')
        ss('Eloqua voice recording screen showing the record button and sentence to read', h=3.8)

        step_heading(9, 'Complete your first voice check')
        body('Eloqua will guide you through a short voice assessment.')
        body('Just follow the instructions on screen. It takes about 5 minutes.')
        body('This gives Eloqua a starting score for your voice so it can track your progress.')
        tip('Sit comfortably. You can rest between tasks. There is no time pressure.')
        ss('Eloqua voice assessment screen', h=3.8)

        step_heading(10, 'You are ready!')
        body('After the voice check, your home screen will appear.')
        body('Your Eloqua journey has started.')
        body('Try to complete one voice session every day for the best results.')
        ss('Eloqua home screen showing the training roadmap', h=3.8)

    # Android steps
    else:

        step_heading(1, 'Open the Google Play Store')
        body('Find the Play Store icon on your phone.')
        body('It looks like a colourful triangle pointing to the right.')
        body('Tap it to open.')
        tip('The Play Store is usually on your home screen or in your app list.')
        ss('Google Play Store home screen', h=3.8)

        step_heading(2, 'Search for Eloqua')
        body('Tap the search bar at the top of the screen.')
        body('Type  Eloqua  and tap the search key on your keyboard.')
        body('Look for the Eloqua app in the results. It has a teal dolphin icon.')
        body('Tap on it.')
        ss('Play Store search results showing the Eloqua app', h=3.8)

        step_heading(3, 'Install the app')
        body('Tap  Install.')
        body('Eloqua will download to your phone. This usually takes less than a minute.')
        body('When it is done, tap  Open.')
        tip('You can also find the Eloqua icon in your app list and tap it there.')
        ss('Eloqua page in the Play Store with the Install button', h=3.8)

        step_heading(4, 'Create your account')
        body('Tap  Sign Up.')
        body('Enter your email address.')
        body('Choose a password. Use at least 8 characters.')
        body('Tap  Create Account.')
        tip('Write your email and password down somewhere safe before continuing.')
        ss('Eloqua Sign Up screen', h=3.8)

        step_heading(5, 'Allow microphone access')
        body('A message will appear asking if Eloqua can use your microphone.')
        body('Tap  Allow.')
        body('The microphone is needed so Eloqua can hear your voice during exercises.')
        tip(
            'If you tapped  Deny  by mistake, go to:\n'
            'Settings  >  Apps  >  Eloqua  >  Permissions\n'
            'Turn on Microphone.'
        )
        ss('Android microphone permission popup', h=2.5)

        step_heading(6, 'Tell us about yourself')
        body('Type your first name.')
        body('Enter your age.')
        body('Tap  Continue.')
        ss('Eloqua About You screen', h=3.8)

        step_heading(7, 'Record your voice  (3 short sentences)')
        body('You will be asked to read 3 sentences out loud.')
        body('This helps Eloqua learn to recognise your voice.')
        body('Tap the record button, read the sentence clearly, then tap stop.')
        body('Repeat this for all 3 sentences.')
        tip('Find a quiet room. Hold your phone about 20 cm (8 inches) from your mouth.')
        ss('Eloqua voice recording screen showing the record button and sentence to read', h=3.8)

        step_heading(8, 'Complete your first voice check')
        body('Eloqua will guide you through a short voice assessment.')
        body('Just follow the instructions on screen. It takes about 5 minutes.')
        body('This gives Eloqua a starting score for your voice so it can track your progress.')
        tip('Sit comfortably. You can rest between tasks. There is no time pressure.')
        ss('Eloqua voice assessment screen', h=3.8)

        step_heading(9, 'You are ready!')
        body('After the voice check, your home screen will appear.')
        body('Your Eloqua journey has started.')
        body('Try to complete one voice session every day for the best results.')
        ss('Eloqua home screen showing the training roadmap', h=3.8)

    # Common problems
    doc.add_paragraph()
    add_rule(doc)
    doc.add_paragraph()

    h2 = doc.add_paragraph()
    h2r = h2.add_run('Common problems and fixes')
    h2r.bold = True
    h2r.font.size = Pt(16)
    h2r.font.color.rgb = RGBColor(0x2B, 0x7A, 0x78)
    h2.paragraph_format.space_before = Pt(4)

    problems = (
        [
            ('I cannot find the invitation email.',
             'Check your Junk or Spam folder. Search your inbox for the word "TestFlight".'),
            ('The app will not open.',
             'Turn your iPhone off and on again, then try opening Eloqua.'),
            ('Eloqua cannot hear my voice.',
             'Go to Settings > Privacy & Security > Microphone, find Eloqua, and turn it on.'),
            ('The app says "Server error" or "Try again later".',
             'Check you are connected to Wi-Fi, then wait 30 seconds and try again.'),
            ('I forgot my password.',
             'On the Sign In screen, tap "Forgot password?" and follow the steps. You will get an email.'),
        ] if is_ios else [
            ('I cannot find Eloqua in the Play Store.',
             'Make sure your phone is running Android 10 or newer. Try searching "Eloqua voice".'),
            ('The app will not open.',
             'Turn your phone off and on again, then try opening Eloqua.'),
            ('Eloqua cannot hear my voice.',
             'Go to Settings > Apps > Eloqua > Permissions and turn on Microphone.'),
            ('The app says "Server error" or "Try again later".',
             'Check you are connected to Wi-Fi, then wait 30 seconds and try again.'),
            ('I forgot my password.',
             'On the Sign In screen, tap "Forgot password?" and follow the steps. You will get an email.'),
        ]
    )

    for problem, fix in problems:
        qp = doc.add_paragraph()
        qr = qp.add_run('Q:  ' + problem)
        qr.bold = True
        qr.font.size = Pt(13)
        qp.paragraph_format.space_before = Pt(8)
        qp.paragraph_format.space_after = Pt(1)

        ap = doc.add_paragraph()
        ar = ap.add_run('A:  ' + fix)
        ar.font.size = Pt(13)
        ap.paragraph_format.left_indent = Inches(0.3)
        ap.paragraph_format.space_before = Pt(1)
        ap.paragraph_format.space_after = Pt(4)

    # Support
    doc.add_paragraph()
    add_rule(doc)
    doc.add_paragraph()

    h3 = doc.add_paragraph()
    h3r = h3.add_run('Need more help?')
    h3r.bold = True
    h3r.font.size = Pt(16)
    h3r.font.color.rgb = RGBColor(0x2B, 0x7A, 0x78)

    add_callout(
        doc,
        'Email us at:  devansh0505@gmail.com\n\n'
        'Please include:\n'
        '  -  Your name\n'
        '  -  The type of phone you are using\n'
        '  -  What happened and which step you were on\n\n'
        'We aim to reply within 24 hours.',
        bg_hex='E8F5F4'
    )

    doc.add_paragraph()

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run('Eloqua  -  Voice Training App  -  2026')
    fr.font.size = Pt(10)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    return doc


if __name__ == '__main__':
    import sys
    out = r'C:\Users\devan\Desktop\Imperial\Eloqua\Eloquav2\Eloquav2'
    suffix = '_new' if '--new' in sys.argv else ''
    build_doc('ios').save(out + r'\Eloqua_Setup_Guide_iOS' + suffix + '.docx')
    print('iOS saved')
    build_doc('android').save(out + r'\Eloqua_Setup_Guide_Android' + suffix + '.docx')
    print('Android saved')
